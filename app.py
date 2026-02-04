from flask import Flask, request, jsonify, render_template, send_file
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
from docx import Document
from pypdf import PdfReader
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
from openai import OpenAI

import os, json, io, datetime, random

app = Flask(__name__)

# ================= OPEN QUESTIONS (FLAN) =================

tokenizer = AutoTokenizer.from_pretrained("google/flan-t5-small")
flan = AutoModelForSeq2SeqLM.from_pretrained("google/flan-t5-small")

def generate_open_questions(text, count):
    text = text[:800]
    questions = []
    attempts = 0
    max_attempts = count * 5   # защита от зависания

    while len(questions) < count and attempts < max_attempts:
        attempts += 1

        prompt = f"""
Generate ONE clear exam question based on the text.
The question MUST end with a question mark.
Do NOT repeat previous questions.
Return ONLY the question.

TEXT:
{text}
"""

        inputs = tokenizer(
            prompt,
            return_tensors="pt",
            truncation=True,
            max_length=512
        )

        out = flan.generate(
            **inputs,
            max_length=128,
            do_sample=True,
            temperature=1.1,
            top_p=0.95
        )
def generate_open_questions(text, count):
    # FLAN физически слабый, поэтому ограничиваем вход
    text = text[:700]

    questions = []
    attempts = 0
    max_attempts = count * 4   # защита от вечного цикла

    # --- 1. ПЫТАЕМСЯ ВЫЖАТЬ ИЗ FLAN РЕАЛЬНЫЕ ВОПРОСЫ ---
    while len(questions) < count and attempts < max_attempts:
        attempts += 1

        prompt = (
            "Generate ONE clear exam question based on the text.\n"
            "The question MUST end with a question mark.\n"
            "Do NOT repeat previous questions.\n"
            "Return ONLY the question.\n\n"
            f"TEXT:\n{text}"
        )

        inputs = tokenizer(
            prompt,
            return_tensors="pt",
            truncation=True,
            max_length=512
        )

        output = flan.generate(
            **inputs,
            max_length=80,
            do_sample=True,
            temperature=1.1,
            top_p=0.9
        )

        q = tokenizer.decode(output[0], skip_special_tokens=True).strip()

        # фильтр адекватности
        if (
            q.endswith("?")
            and len(q) > 15
            and q not in questions
        ):
            questions.append(q)

    # --- 2. УМНЫЙ FALLBACK (РАЗНЫЕ ВОПРОСЫ, НЕ КОПИПАСТА) ---
    fallback_templates = [
        "What is the main idea discussed in the text?",
        "Which concept in the text is the most important and why?",
        "How does the text explain its central theme?",
        "What key argument is presented in the text?",
        "Which idea from the text is most significant for understanding the topic?",
        "What process or phenomenon is described in the text?",
        "How does the text relate its main concept to the broader topic?"
    ]

    i = 0
    while len(questions) < count:
        questions.append(fallback_templates[i % len(fallback_templates)])
        i += 1

    return questions




# ================= MCQ (GPT) =================

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

def shuffle_mcq(q):
    pairs = list(enumerate(q["options"]))
    random.shuffle(pairs)

    new_opts, new_ans = [], []
    for new_i, (old_i, opt) in enumerate(pairs):
        new_opts.append(opt)
        if old_i in q["answer_indices"]:
            new_ans.append(new_i)

    q["options"] = new_opts
    q["answer_indices"] = new_ans
    return q

def generate_mcq(text, count):
    prompt = f"""
Return ONLY valid JSON:
{{
  "questions": [
    {{
      "question": "",
      "options": ["", "", "", ""],
      "answer_indices": [],
      "explanation": ""
    }}
  ]
}}

Rules:
- Exactly {count} questions
- 1 to 3 correct answers
- Correct answers must be randomly distributed
- options WITHOUT letters
- Use ONLY the text below

TEXT:
{text}
"""
    r = client.responses.create(model="gpt-4.1-mini", input=prompt)
    raw = r.output_text
    raw = raw[raw.find("{"):raw.rfind("}")+1]

    qs = json.loads(raw)["questions"]
    return [shuffle_mcq(q) for q in qs]

# ================= FILE TEXT =================

def pdf_text(f):
    return "".join(p.extract_text() or "" for p in PdfReader(f).pages)

def docx_text(f):
    d = Document(f)
    return "\n".join(p.text for p in d.paragraphs)

# ================= HISTORY =================

def save_history(t, q):
    data = []
    if os.path.exists("history.json"):
        data = json.load(open("history.json", encoding="utf-8"))

    data.append({
        "type": t,
        "questions": q,
        "time": str(datetime.datetime.now())
    })

    json.dump(data, open("history.json", "w", encoding="utf-8"),
              ensure_ascii=False, indent=2)

# ================= ROUTES =================

@app.route("/")
def intro():
    return render_template("intro.html")

@app.route("/generator")
def generator():
    return render_template("index.html")

@app.route("/about")
def about():
    return render_template("about.html")

@app.route("/history_page")
def history_page():
    data = []
    if os.path.exists("history.json"):
        data = json.load(open("history.json", encoding="utf-8"))
    return render_template("history.html", history=data)

@app.route("/exam/<int:exam_id>")
def view_exam(exam_id):
    data = json.load(open("history.json", encoding="utf-8"))
    return render_template("exam.html", exam=data[exam_id])

# ================= API =================

@app.route("/upload_file", methods=["POST"])
def upload_file():
    f = request.files["file"]
    if f.filename.endswith(".pdf"):
        return jsonify({"text": pdf_text(f)})
    if f.filename.endswith(".docx"):
        return jsonify({"text": docx_text(f)})
    return jsonify({"error": "bad file"}), 400

@app.route("/generate_open", methods=["POST"])
def gen_open():
    d = request.json
    q = generate_open_questions(d["text"], int(d["count"]))
    save_history("open", q)
    return jsonify({"questions": q})

@app.route("/generate_mcq", methods=["POST"])
def gen_mcq():
    d = request.json
    q = generate_mcq(d["text"], int(d["count"]))
    save_history("mcq", q)
    return jsonify({"questions": q})

# ================= EXPORT =================

@app.route("/download_docx", methods=["POST"])
def download_docx():
    qs = request.json["questions"]
    d = Document()

    for i, q in enumerate(qs, 1):
        if isinstance(q, dict):
            d.add_paragraph(f"{i}. {q['question']}")
            for j, o in enumerate(q["options"]):
                d.add_paragraph(f"{chr(65+j)}. {o}")
            ans = ", ".join(chr(65+i) for i in q["answer_indices"])
            d.add_paragraph(f"Correct answers: {ans}")
            d.add_paragraph(q["explanation"])
        else:
            d.add_paragraph(f"{i}. {q}")

    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name="exam.docx")

@app.route("/download_pdf", methods=["POST"])
def download_pdf():
    qs = request.json["questions"]
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    story = []

    for i, q in enumerate(qs, 1):
        if isinstance(q, dict):
            story.append(Paragraph(f"<b>{i}. {q['question']}</b>", styles["Normal"]))
            for j, o in enumerate(q["options"]):
                story.append(Paragraph(f"{chr(65+j)}. {o}", styles["Normal"]))
            ans = ", ".join(chr(65+i) for i in q["answer_indices"])
            story.append(Paragraph(f"<b>Correct answers:</b> {ans}", styles["Normal"]))
            story.append(Paragraph(q["explanation"], styles["Normal"]))
        else:
            story.append(Paragraph(f"<b>{i}. {q}</b>", styles["Normal"]))
        story.append(Spacer(1, 12))

    doc.build(story)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name="exam.pdf")

if __name__ == "__main__":
    app.run(debug=True)
